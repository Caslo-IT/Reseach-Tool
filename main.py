import json
import os
import re
import time
from datetime import datetime
from xml.sax.saxutils import escape

from google import genai
from google.genai import errors as genai_errors
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from unicode_to_font import convertor

SINHALA_FONT_NAME = "FMGanganee"
SINHALA_FONT_FILE = os.path.join("fonts", "FM Ganganee x.ttf")


def load_dotenv():
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if not os.path.exists(env_path):
        return

    with open(env_path, encoding="utf-8") as env_file:
        for raw_line in env_file:
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip("'\""))


load_dotenv()

MODEL_NAME = "gemma-4-31b-it"
MAX_GEMINI_RETRIES = 3

prompt_template = """You are an expert researcher in Sri Lankan sacred art, Hindu-Buddhist iconography, temple sculpture traditions, and statue design interpretation.

I will provide a statue / deity / sacred figure description or research topic. Your job is to prepare a deeply researched, long-form report suitable for a 15 to 20 page PDF.

IMPORTANT REQUIREMENTS:
- Treat the user input as the research subject.
- Produce a detailed, comprehensive, professional report with substantial explanation under every required section.
- Write the full report primarily in Sinhala.
- Use English only for unavoidable technical terms, and when used, keep them minimal and preferably in brackets.
- Include multiple useful tables throughout the report.
- At the end of the report, include a resource section with direct source locations.
- Include webpage links and YouTube video links when relevant.
- If a detail is unknown or not directly inferable, write `සඳහන් නොවේ` or `තහවුරු කිරීම අවශ්‍යයි`.
- Do not skip any required sections.
- Expand each area with interpretive, symbolic, artistic, cultural, and practical design analysis.
- The output should be useful for statue design workflow, 3D modeling reference, and religious/iconographic validation.

OUTPUT FORMAT RULES:
- Return clean markdown-style content only.
- Use `#` for the main title.
- Use `##` for main sections.
- Use `###` for subsections.
- Use `-` for bullet points.
- Use markdown tables with `|` separators where appropriate.
- Do not wrap the answer in code fences.
- All section titles, explanations, bullets, and table text must be in Sinhala as much as possible.
- In the final resource section, include direct URLs in plain text so they can be copied easily.

REQUIRED REPORT STRUCTURE:
# පර්යේෂණ වාර්තා ශීර්ෂය

## සංක්ෂිප්ත සාරාංශය
- විෂයය, ආලේඛනමය වැදගත්කම, කලාත්මක අනන්‍යතාව, සහ නිර්මාණමය වැදගත්කම සාරාංශ කරන්න.

## හැඳින්වීම
- ප්‍රතිමාව / දෙවියන් / පූජනීය චරිතය කවුද හෝ කුමක්දැයි පැහැදිලි කරන්න.
- ආගමික, සංස්කෘතික, සහ කලාත්මක පසුබිම පැහැදිලි කරන්න.
- ශ්‍රී ලාංකේය පූජනීය කලා සම්ප්‍රදායේ හෝ ප්‍රතිමා නිර්මාණයේ මේ විෂයය වැදගත් වන්නේ ඇයිදැයි සඳහන් කරන්න.

## 1. අංගලක්ෂණ සහ සමානුපාතික මිනුම් (Physical Attributes & Proportions)
### 1.1 තාළමාන ක්‍රමය සහ අවකාශීය මිනුම් න්‍යාය
- සමානුපාතිකත්වය, මිනුම් ක්‍රමය, ශරීර සමබරතාව, ඉරියව් තර්කය, සහ අවකාශීය සංයෝජනය විස්තර කරන්න.

### 1.2 උත්තම දශතාළ මිනුම් පද්ධතිය
- හඳුනාගත හැකි නම් තාළමාන ක්‍රමය සඳහන් කරන්න.
- නිශ්චිත තාළමාන ක්‍රමය පැහැදිලි නොවේ නම්, සම්ප්‍රදායික මිනුම් ප්‍රවේශය පැහැදිලි කරන්න.

### 1.3 මුහුණේ ඉරියව්, පෞරුෂය සහ දිව්‍යමය ස්වභාවය
- මුහුණේ ප්‍රකාශනය, දිව්‍යමය භාවය, හැඟීම්මය ස්වභාවය, පෞරුෂය, සහ පූජනීය දෘශ්‍ය අනන්‍යතාව විස්තර කරන්න.

## 2. ආභරණ සහ සළුපිළි (Clothing & Ornaments)
- කිරුළ, ආභරණ, වස්ත්‍ර, ශරීර අලංකාර, සංකේතාත්මක ද්‍රව්‍ය, අලංකාර රටාව, සහ ශෛලීය සම්ප්‍රදාය විස්තර කරන්න.
- සළුපිළි හා ආභරණ පිලිදව වගුවක් අශ්‍රයෙන් විස්තර ඉදිරිපත් කරන්න

## 3. අත්වල ඉරියව් සහ ආයුධ/උපකරණ (Hand Gestures & Weapons)
### 3.1 මුද්‍රා සහ ඇඟිලි පිහිටීම (Mudras)
- අත් ඉරියව්, ඇඟිලි පිහිටීම, පූජාත්මක අර්ථය, සංකේතාත්මක අභිප්‍රාය, සහ දෘශ්‍ය සන්නිවේදනය විස්තර කරන්න.

### 3.2 ආයුධ සහ උපකරණ (Ayudhas)
- ආයුධ, අතේ දරා සිටින වස්තු, පූජා උපකරණ, සංකේතාත්මක ද්‍රව්‍ය, සහ ඒවායේ අර්ථය පැහැදිලි කරන්න.
- ආයුධ, අතේ දරා සිටින වස්තු, පූජා උපකරණ, සංකේතාත්මක ද්‍රව්‍ය පිලිදව වගුවක් අශ්‍රයෙන් විස්තර ඉදිරිපත් කරන්න


## 4. පාදම හෙවත් ආසනය (Pedestal/Base)
### 4.1 පද්ම පීඨය / භද්‍ර පීඨය
- පාදමේ වර්ගය, සංකේතාත්මක අර්ථය, ව්‍යුහමය වැදගත්කම, සහ නිර්මාණමය අර්ථකථනය විස්තර කරන්න.

## 5. පුරාවෘත්ත, සම්බන්ධතා සහ සලකුණු (Mythology, Relationships & Symbols)
### 5.1 සම්බන්දතා
- සම්බන්ධ දෙවිවරු, දිව්‍ය පවුල් සම්බන්ධතා, කතානායක සබැඳි, සහ පුරාවෘත්තය පැහැදිලි කරන්න.

### 5.2 වාහනය
- වාහනයක් තිබේ නම් එය සහ එහි සංකේතාත්මක අර්ථය පැහැදිලි කරන්න.

### 5.3 ශ්‍රී ලාංකේය සම්ප්‍රදාය සහ දාර්ශනික පසුබිම
- ශ්‍රී ලාංකේය සංස්කෘතික අර්ථකථනය, ප්‍රාදේශීය භක්ති සම්ප්‍රදාය, සහ දාර්ශනික පසුබිම විස්තර කරන්න.

## සංසන්දනාත්මක සහ නිර්මාණමය අර්ථකථනය
- නිරූපණයේ විය හැකි වෙනස්කම් සංසන්දනය කරන්න.
- ශාස්ත්‍රීය, ප්‍රාදේශීය, සහ නවීන දෘශ්‍ය ප්‍රවේශවල වෙනස්කම් සාකච්ඡා කරන්න.
- එම විස්තර ප්‍රතිමා නිර්මාණ තීරණවලට බලපාන ආකාරය පැහැදිලි කරන්න.

## ප්‍රායෝගික නිර්මාණ මාර්ගෝපදේශ
- මූර්ති ශිල්පීන්, නිර්මාණකරුවන්, හෝ 3D ආකෘති නිර්මාණකරුවන් සඳහා ප්‍රයෝජනවත් මාර්ගෝපදේශ ලබා දෙන්න.
- ආලේඛනමය නිවැරදිභාවය සඳහා රැකගත යුතු කරුණු සඳහන් කරන්න.
- නිෂ්පාදනයට පෙර විශේෂයෙන් තහවුරු කළ යුතු අංශ සඳහන් කරන්න.

## සාරාංශ වගු
- පහත වැනි වගු කිහිපයක් ඇතුළත් කරන්න:
- ශාරීරික අංගලක්ෂණ සාරාංශ වගුව
- ආභරණ සහ වස්ත්‍ර සාරාංශ වගුව
- මුද්‍රා / ආයුධ අර්ථ වගුව
- සංකේත සහ පුරාවෘත්ත සාරාංශ වගුව
- නිර්මාණ සත්‍යාපන පිරික්සුම් ලැයිස්තුව

## නිගමනය
- ප්‍රතිමාව / දෙවියන් / පූජනීය විෂයය පිළිබඳ ශක්තිමත් අවසන් අර්ථකථනයක් ලබා දෙන්න.

## මූලාශ්‍ර සහ සම්පත් ස්ථාන
- අවසානයේ භාවිත කළ හෝ යෝජනා කරන මූලාශ්‍ර ලැයිස්තුවක් ලබා දෙන්න.
- අවම වශයෙන් පහත තොරතුරු ඇතුළත් කරන්න:
- වෙබ් පිටු සබැඳි (Web pages)
- YouTube වීඩියෝ සබැඳි
- හැකි නම් එක් එක් සබැඳියට කෙටි විස්තරයක්
- සබැඳි markdown වගුවක් හෝ bullet list එකක් ලෙස දෙන්න.
- URL සම්පූර්ණ ආකාරයෙන් `https://...` ලෙස පෙන්වන්න.

USER RESEARCH SUBJECT:
{user_description}
"""

sketch_prompt_template = """You are an expert prompt writer for AI image generation.

Use the following research report content to create image-generation prompts for sketch images.

STRICT REQUIREMENTS:
- Return JSON array only.
- Do not return markdown.
- Do not return explanation text.
- Generate exactly 30 objects.
- IDs must run from 1 to 30 in order.
- Items 1 to 10 must be full-body prompts.
- Items 11 to 30 must be close-up prompts.
- Every prompt must clearly specify a square 1:1 composition.
- Every prompt must describe a high-detail pencil or graphite sketch on a clean white paper background.
- Keep the deity / statue design consistent with the research report.
- Make every title unique.
- Each object must follow this exact schema:
  {{
    "id": 1,
    "title": "Title here",
    "prompt": "Prompt here"
  }}
- Prompt wording must be production-ready and visually descriptive.
- Mention whether the image is full-body or close-up naturally inside each prompt.
- Emphasize clean linework, refined shading, and sharp detailing.

RESEARCH REPORT CONTENT:
{research_text}
"""


def build_client():
    api_key = os.getenv("GEMINI_API_KEY", "")
    if not api_key:
        raise ValueError("GEMINI_API_KEY is not set. Add it to .env or your environment.")
    return genai.Client(api_key=api_key)


def _should_retry_gemini_error(error):
    status_code = getattr(error, "status_code", None)
    if isinstance(error, genai_errors.ServerError):
        return True
    return isinstance(status_code, int) and status_code >= 500


def generate_content_with_retry(client, prompt, *, progress_callback=None, stage_label="Generating content"):
    last_error = None

    for attempt in range(1, MAX_GEMINI_RETRIES + 1):
        try:
            return client.models.generate_content(
                model=MODEL_NAME,
                contents=prompt,
            )
        except genai_errors.APIError as error:
            last_error = error
            if not _should_retry_gemini_error(error) or attempt == MAX_GEMINI_RETRIES:
                break

            wait_seconds = attempt * 2
            if progress_callback:
                progress_callback(
                    min(0.95, 0.12 + (attempt * 0.02)),
                    f"{stage_label} failed temporarily. Retrying in {wait_seconds} seconds (attempt {attempt + 1}/{MAX_GEMINI_RETRIES})",
                )
            time.sleep(wait_seconds)

    if last_error:
        raise RuntimeError(
            "Gemini is temporarily unavailable and the research could not be generated right now. "
            "Please try again in a moment."
        ) from last_error

    raise RuntimeError("Gemini generation failed before a response was returned.")


def _read_usage_value(usage, *keys):
    for key in keys:
        if isinstance(usage, dict) and key in usage:
            value = usage.get(key)
        else:
            value = getattr(usage, key, None)
        if isinstance(value, int):
            return value
    return 0


def extract_usage_metrics(response):
    usage = getattr(response, "usage_metadata", None) or getattr(response, "usageMetadata", None)
    if not usage:
        return {
            "input_tokens": 0,
            "output_tokens": 0,
            "total_tokens": 0,
        }

    input_tokens = _read_usage_value(usage, "prompt_token_count", "promptTokenCount")
    output_tokens = _read_usage_value(usage, "candidates_token_count", "candidatesTokenCount")
    total_tokens = _read_usage_value(usage, "total_token_count", "totalTokenCount")

    if total_tokens == 0:
        total_tokens = input_tokens + output_tokens

    return {
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": total_tokens,
    }


def combine_usage_metrics(*metrics):
    return {
        "input_tokens": sum(metric.get("input_tokens", 0) for metric in metrics),
        "output_tokens": sum(metric.get("output_tokens", 0) for metric in metrics),
        "total_tokens": sum(metric.get("total_tokens", 0) for metric in metrics),
    }


def build_paragraph_text(text):
    if not text:
        return "&nbsp;"

    parts = re.split(r"([\u0D80-\u0DFF]+)", text)
    formatted_parts = []

    for part in parts:
        if not part:
            continue
        if re.fullmatch(r"[\u0D80-\u0DFF]+", part):
            formatted_parts.append(
                f'<font name="{SINHALA_FONT_NAME}">{escape(convertor(part, "fm"))}</font>'
            )
        else:
            formatted_parts.append(escape(part))

    return "".join(formatted_parts) or "&nbsp;"


def parse_table_row(line):
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return None
    return [cell.strip() for cell in stripped.strip("|").split("|")]


def is_separator_row(cells):
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def build_table(table_lines, body_style):
    rows = []
    for line in table_lines:
        cells = parse_table_row(line)
        if cells is not None:
            rows.append(cells)

    if not rows:
        return []

    if len(rows) > 1 and is_separator_row(rows[1]):
        rows.pop(1)

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]

    cell_data = []
    for row in normalized_rows:
        cell_data.append([Paragraph(build_paragraph_text(cell), body_style) for cell in row])

    table = Table(cell_data, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E6F2")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#7A8A99")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 6),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
            ]
        )
    )
    return [table, Spacer(1, 0.15 * inch)]


def build_story(markdown_text, styles):
    story = []
    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            story.append(Spacer(1, 0.12 * inch))
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                else:
                    break
            story.extend(build_table(table_lines, styles["Body"]))
            continue

        if stripped.startswith("# "):
            story.append(Paragraph(build_paragraph_text(stripped[2:].strip()), styles["Title"]))
            story.append(Spacer(1, 0.18 * inch))
            i += 1
            continue

        if stripped.startswith("## "):
            story.append(Paragraph(build_paragraph_text(stripped[3:].strip()), styles["Heading1"]))
            story.append(Spacer(1, 0.1 * inch))
            i += 1
            continue

        if stripped.startswith("### "):
            story.append(Paragraph(build_paragraph_text(stripped[4:].strip()), styles["Heading2"]))
            story.append(Spacer(1, 0.08 * inch))
            i += 1
            continue

        if stripped.startswith("- "):
            bullet_text = build_paragraph_text(stripped[2:].strip())
            story.append(Paragraph(bullet_text, styles["Bullet"]))
            i += 1
            continue

        story.append(Paragraph(build_paragraph_text(stripped), styles["Body"]))
        story.append(Spacer(1, 0.06 * inch))
        i += 1

    return story


def register_fonts():
    font_path = os.path.join(os.path.dirname(__file__), SINHALA_FONT_FILE)
    if not os.path.exists(font_path):
        raise FileNotFoundError(f"Sinhala font not found: {font_path}")
    pdfmetrics.registerFont(TTFont(SINHALA_FONT_NAME, font_path))


def set_run_font(run, font_name, font_size=None, bold=False):
    run.bold = bold
    run.font.name = font_name
    if font_size is not None:
        run.font.size = font_size

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font_name)


def add_docx_text(paragraph, text, font_name, font_size=None, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size=font_size, bold=bold)
    return run


def style_docx_table(table):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, "Nirmala UI", font_size=Pt(10.5), bold=(row_index == 0))


def build_docx(markdown_text, output_path):
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    lines = markdown_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                else:
                    break

            rows = []
            for table_line in table_lines:
                cells = parse_table_row(table_line)
                if cells is not None:
                    rows.append(cells)

            if len(rows) > 1 and is_separator_row(rows[1]):
                rows.pop(1)

            if rows:
                max_cols = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=max_cols)
                for row_index, row in enumerate(rows):
                    for col_index in range(max_cols):
                        cell = table.cell(row_index, col_index)
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        add_docx_text(
                            paragraph,
                            row[col_index] if col_index < len(row) else "",
                            "Nirmala UI",
                            font_size=Pt(10.5),
                            bold=(row_index == 0),
                        )
                style_docx_table(table)
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            add_docx_text(paragraph, stripped[2:].strip(), "Nirmala UI", font_size=Pt(18), bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph()
            add_docx_text(paragraph, stripped[3:].strip(), "Nirmala UI", font_size=Pt(14), bold=True)
            i += 1
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph()
            add_docx_text(paragraph, stripped[4:].strip(), "Nirmala UI", font_size=Pt(12), bold=True)
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_docx_text(paragraph, stripped[2:].strip(), "Nirmala UI", font_size=Pt(10.5))
            i += 1
            continue

        paragraph = document.add_paragraph()
        add_docx_text(paragraph, stripped, "Nirmala UI", font_size=Pt(10.5))
        i += 1

    document.save(output_path)


def build_styles():
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#183A5A"),
            spaceAfter=10,
        ),
        "Heading1": ParagraphStyle(
            "ReportHeading1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#204C73"),
            spaceBefore=8,
            spaceAfter=6,
        ),
        "Heading2": ParagraphStyle(
            "ReportHeading2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12.5,
            leading=16,
            textColor=colors.HexColor("#2B5D87"),
            spaceBefore=6,
            spaceAfter=4,
        ),
        "Body": ParagraphStyle(
            "ReportBody",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.black,
        ),
        "Bullet": ParagraphStyle(
            "ReportBullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            leftIndent=18,
            firstLineIndent=-8,
            bulletIndent=6,
        ),
    }


def build_output_paths():
    resources_dir = os.path.join(os.path.dirname(__file__), "Resources")
    os.makedirs(resources_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_dir = os.path.join(resources_dir, f"report_{timestamp}")
    os.makedirs(report_dir, exist_ok=True)

    return {
        "report_dir": report_dir,
        "pdf_path": os.path.join(report_dir, "research_report.pdf"),
        "docx_path": os.path.join(report_dir, "research_report.docx"),
        "json_path": os.path.join(report_dir, "sketch_prompts.json"),
    }


def extract_json_array(raw_text):
    text = raw_text.strip()

    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()

    data = json.loads(text)
    if not isinstance(data, list):
        raise ValueError("Sketch prompt response was not a JSON array.")
    if len(data) != 30:
        raise ValueError(f"Expected 30 sketch prompts, but received {len(data)}.")
    return data


def validate_sketch_prompts(items):
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"Sketch prompt item {index} is not an object.")
        if item.get("id") != index:
            raise ValueError(f"Sketch prompt item {index} has invalid id {item.get('id')}.")
        if not isinstance(item.get("title"), str) or not item["title"].strip():
            raise ValueError(f"Sketch prompt item {index} is missing a valid title.")
        if not isinstance(item.get("prompt"), str) or not item["prompt"].strip():
            raise ValueError(f"Sketch prompt item {index} is missing a valid prompt.")


def generate_sketch_prompts(client, research_text):
    response = generate_content_with_retry(
        client,
        sketch_prompt_template.format(research_text=research_text),
        stage_label="Generating sketch prompts",
    )
    raw_text = response.text or "[]"
    items = extract_json_array(raw_text)
    validate_sketch_prompts(items)
    return items, extract_usage_metrics(response)


def generate_report_assets(user_description, progress_callback=None):
    if not user_description or not user_description.strip():
        raise ValueError("Research subject input is required.")

    started_at = time.perf_counter()

    def update_progress(fraction, message):
        if progress_callback:
            progress_callback(fraction, message)

    update_progress(0.05, "Initializing Gemini client")
    client = build_client()

    update_progress(0.15, "Generating research report content")
    prompt = prompt_template.format(user_description=user_description.strip())
    response = generate_content_with_retry(
        client,
        prompt,
        progress_callback=update_progress,
        stage_label="Generating research report",
    )
    text = response.text or "No response text was returned."
    report_usage = extract_usage_metrics(response)

    update_progress(0.45, "Preparing PDF layout and output paths")
    register_fonts()
    styles = build_styles()
    output_paths = build_output_paths()

    update_progress(0.6, "Generating sketch prompt JSON")
    sketch_prompts, sketch_usage = generate_sketch_prompts(client, text)

    update_progress(0.78, "Building research PDF")
    doc = SimpleDocTemplate(
        output_paths["pdf_path"],
        pagesize=A4,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
    )
    story = build_story(text, styles)
    doc.build(story)

    update_progress(0.88, "Building research Word document")
    build_docx(text, output_paths["docx_path"])

    update_progress(0.94, "Saving sketch prompt JSON")
    with open(output_paths["json_path"], "w", encoding="utf-8") as json_file:
        json.dump(sketch_prompts, json_file, ensure_ascii=False, indent=2)

    update_progress(1.0, "Completed")
    duration_seconds = round(time.perf_counter() - started_at, 2)
    usage = combine_usage_metrics(report_usage, sketch_usage)
    return {
        "report_text": text,
        "sketch_prompts": sketch_prompts,
        "metrics": {
            "requests": 2,
            "generation_time_seconds": duration_seconds,
            **usage,
        },
        **output_paths,
    }


def main():
    user_description = input("Enter the statue / deity research subject to send to Gemini:\n").strip()
    result = generate_report_assets(user_description)

    print(f"PDF Generated: {result['pdf_path']}")
    print(f"Word Generated: {result['docx_path']}")
    print(f"JSON Generated: {result['json_path']}")


if __name__ == "__main__":
    main()
